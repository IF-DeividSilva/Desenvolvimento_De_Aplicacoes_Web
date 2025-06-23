import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import schemas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY


def gerar_docx_lista_exercicios(lista: schemas.ListaExercicios):
    document = Document()
    titulo = document.add_heading(lista.titulo, level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Nível: {lista.nivel_dificuldade}", style='Intense Quote')
    document.add_paragraph()

    gabarito = []
    for i, exercicio in enumerate(lista.exercicios, start=1):
        enunciado_p = document.add_paragraph()
        enunciado_p.add_run(f"{i}. {exercicio.enunciado}").bold = True
        if exercicio.opcoes:
            for key, value in exercicio.opcoes.items():
                document.add_paragraph(f"   {key}) {value}", style='List Bullet')
        document.add_paragraph()
        gabarito.append(f"{i}. {exercicio.resposta_correta.upper()}")

    document.add_page_break()
    document.add_heading('Gabarito', level=2)
    for resposta in gabarito:
        document.add_paragraph(resposta)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def gerar_pdf_lista_exercicios(lista: schemas.ListaExercicios):
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(lista.titulo, styles['h1']))
    story.append(Paragraph(f"Nível: {lista.nivel_dificuldade}", styles['h2']))
    story.append(Spacer(1, 12))
    
    gabarito = []
    for i, exercicio in enumerate(lista.exercicios, start=1):
        story.append(Paragraph(f"<b>{i}. {exercicio.enunciado}</b>", styles['Normal']))
        story.append(Spacer(1, 12))
        if exercicio.opcoes:
            for key, value in exercicio.opcoes.items():
                story.append(Paragraph(f"   {key}) {value}", styles['Normal']))
        story.append(Spacer(1, 24))
        gabarito.append(f"{i}. {exercicio.resposta_correta.upper()}")
        
    story.append(PageBreak())
    story.append(Paragraph("Gabarito", styles['h2']))
    for resposta in gabarito:
        story.append(Paragraph(resposta, styles['Normal']))
        
    doc.build(story)
    file_stream.seek(0)
    return file_stream


def gerar_docx_texto_apoio(texto: schemas.TextoGerado):
    """
    Cria um DOCX processando o Markdown do texto de apoio.
    """
    document = Document()
    titulo = document.add_heading(texto.tema, level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Matéria: {texto.materia}", style='Intense Quote')
    document.add_paragraph(f"Nível: {texto.nivel}", style='Intense Quote')
    document.add_paragraph()
    
    # Processamento básico de markdown - implementação simplificada
    paragrafos = texto.conteudo.split('\n\n')
    for paragrafo in paragrafos:
        if paragrafo.startswith('# '):
            document.add_heading(paragrafo[2:], level=1)
        elif paragrafo.startswith('## '):
            document.add_heading(paragrafo[3:], level=2)
        elif paragrafo.startswith('### '):
            document.add_heading(paragrafo[4:], level=3)
        else:
            document.add_paragraph(paragrafo)
    
    # Salvar em um BytesIO para retornar como arquivo
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def markdown_to_reportlab_html(markdown_text: str) -> str:
    """
    Converte um Markdown simples para tags HTML que o ReportLab entende. (VERSÃO CORRIGIDA)
    """
    html_text = ""
    for line in markdown_text.splitlines():
        line = line.strip()

        # Verifica cada tipo de formatação
        if not line:
            html_text += "<br/><br/>"
        elif line.startswith('## '):
            html_text += f"<h2>{line.lstrip('## ').strip()}</h2>"
        elif line.startswith('# '):
            html_text += f"<h1>{line.lstrip('# ').strip()}</h1>"
        elif line.startswith('* '):
            processed_line = line.lstrip('* ').strip()
            parts = processed_line.split('**')
            final_line = ""
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    final_line += f"<b>{part}</b>"
                else:
                    final_line += part
            html_text += f"&bull; {final_line}<br/>"
        else:
            parts = line.split('**')
            final_line = ""
            for i, part in enumerate(parts):
                if i % 2 == 1: 
                    final_line += f"<b>{part}</b>"
                else:
                    final_line += part
            html_text += final_line
    return html_text

def gerar_pdf_texto_apoio(texto: schemas.TextoGerado):
    """
    Cria um PDF processando o Markdown do texto de apoio.
    """
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=letter)
    styles = getSampleStyleSheet()
    
    styles['h1'].fontSize = 18
    styles['h2'].fontSize = 14
    
    story = [
        Paragraph(texto.tema, styles['h1']),
        Spacer(1, 12),
    ]

    conteudo_html = markdown_to_reportlab_html(texto.conteudo)
    story.append(Paragraph(conteudo_html, styles['BodyText']))
    
    doc.build(story)
    file_stream.seek(0)
    return file_stream